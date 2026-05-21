from __future__ import annotations

import json
import os
import shutil
from pathlib import Path
from typing import Annotated, Any

import duckdb
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from langchain_core.tools import tool
from langchain_experimental.utilities import PythonREPL
from matplotlib import pyplot as plt
from pypdf import PdfReader

from agent_tools.runtime import (
    artifact_path,
    attachment_manifest,
    collect_runtime_artifacts,
    get_tool_runtime_context,
    register_runtime_artifact,
    resolve_runtime_attachment,
)


plt.switch_backend("Agg")


def _snapshot_workspace_files(context) -> set[str]:
    roots = [context.artifact_dir, context.workspace_dir]
    snapshot: set[str] = set()
    for root in roots:
        if not root.exists():
            continue
        for path in root.rglob("*"):
            if path.is_file():
                try:
                    snapshot.add(str(path.relative_to(root)))
                except ValueError:
                    snapshot.add(path.name)
    return snapshot


def _load_dataframe(
    attachment_id: str,
    *,
    sheet_name: str | None = None,
    rows: int | None = None,
) -> pd.DataFrame:
    attachment = resolve_runtime_attachment(attachment_id)
    path = Path(attachment.storage_path)
    if attachment.kind == "csv":
        frame = pd.read_csv(path)
    elif attachment.kind == "spreadsheet":
        frame = pd.read_excel(path, sheet_name=sheet_name or 0)
    elif attachment.kind == "json":
        frame = pd.read_json(path)
    else:
        raise ValueError(f"Attachment {attachment_id} is not a tabular file.")

    if isinstance(frame, dict):
        first_key = next(iter(frame))
        frame = frame[first_key]
    if rows is not None:
        return frame.head(rows)
    return frame


def _extract_document_text_internal(attachment_id: str, max_chars: int = 6000) -> dict[str, Any]:
    attachment = resolve_runtime_attachment(attachment_id)
    path = Path(attachment.storage_path)

    if attachment.kind == "pdf":
        reader = PdfReader(str(path))
        parts: list[str] = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
            if sum(len(part) for part in parts) >= max_chars:
                break
        text = "\n".join(parts)[:max_chars]
        return {
            "attachment_id": attachment_id,
            "file_name": attachment.file_name,
            "kind": attachment.kind,
            "page_count": len(reader.pages),
            "text_excerpt": text,
            "truncated": len(text) >= max_chars,
            "warning": (
                "No extractable text was found. This PDF may be scanned or image-based and may require OCR."
                if not text.strip()
                else None
            ),
        }

    if attachment.kind == "docx":
        document = Document(str(path))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        table_cells = [
            cell.text.strip()
            for table in document.tables
            for row in table.rows
            for cell in row.cells
            if cell.text.strip()
        ]
        text = "\n".join([*paragraphs, *table_cells])[:max_chars]
        return {
            "attachment_id": attachment_id,
            "file_name": attachment.file_name,
            "kind": attachment.kind,
            "paragraph_count": len(paragraphs),
            "table_cell_count": len(table_cells),
            "text_excerpt": text,
            "truncated": len(text) >= max_chars,
        }

    raise ValueError(f"Attachment {attachment_id} is not a supported document file.")


@tool
def inspect_attachments() -> str:
    """Inspect the currently attached files for this turn. Use first before deeper analysis."""
    return json.dumps({"attachments": attachment_manifest()}, ensure_ascii=False, indent=2)


@tool
def preview_tabular_file(
    attachment_id: Annotated[str, "Attachment id returned by inspect_attachments."],
    rows: Annotated[int, "How many preview rows to return."] = 8,
    sheet_name: Annotated[str | None, "Optional worksheet name for .xlsx files."] = None,
) -> str:
    """Preview a CSV/XLSX/JSON attachment. Use this before writing analysis code."""
    attachment = resolve_runtime_attachment(attachment_id)
    frame = _load_dataframe(attachment_id, sheet_name=sheet_name, rows=rows)
    available_sheets: list[str] | None = None
    if attachment.kind == "spreadsheet":
        available_sheets = pd.ExcelFile(attachment.storage_path).sheet_names
    return json.dumps(
        {
            "attachment_id": attachment_id,
            "file_name": attachment.file_name,
            "kind": attachment.kind,
            "sheet_name": sheet_name,
            "available_sheets": available_sheets,
            "columns": list(map(str, frame.columns.tolist())),
            "preview_rows": frame.replace({np.nan: None}).to_dict(orient="records"),
            "row_count_sampled": len(frame),
        },
        ensure_ascii=False,
        indent=2,
        default=str,
    )


@tool
def extract_document_text(
    attachment_id: Annotated[str, "Attachment id returned by inspect_attachments."],
    max_chars: Annotated[int, "Maximum characters to return from the document."] = 6000,
) -> str:
    """Extract text from a PDF or DOCX attachment. Use when the task needs document text.

    Phase 4.3 — corrupted/encrypted/unsupported files surface a structured
    ToolErrorPayload instead of bubbling an opaque library exception, so the
    supervisor/validator can react cleanly and the SSE ``tool_error`` event
    carries an actionable message.
    """
    from agent_tools.errors import make_tool_error_payload

    try:
        result = _extract_document_text_internal(attachment_id, max_chars=max_chars)
    except ValueError as exc:
        # raised when attachment kind is not pdf/docx
        return json.dumps(
            make_tool_error_payload(
                kind="input_validation",
                message=str(exc),
                details={"attachment_id": attachment_id},
            ),
            ensure_ascii=False,
        )
    except FileNotFoundError as exc:
        return json.dumps(
            make_tool_error_payload(
                kind="not_found",
                message=f"document file is missing on disk: {exc}",
                details={"attachment_id": attachment_id},
            ),
            ensure_ascii=False,
        )
    except Exception as exc:
        # pypdf / python-docx raise a variety of subclasses for corrupted or
        # encrypted files; keep the structured envelope so the caller can decide.
        return json.dumps(
            make_tool_error_payload(
                kind="runtime",
                message=f"failed to extract document text: {type(exc).__name__}: {exc}",
                details={"attachment_id": attachment_id},
            ),
            ensure_ascii=False,
        )

    return json.dumps(result, ensure_ascii=False, indent=2)


@tool
def profile_dataframe(
    attachment_id: Annotated[str, "Attachment id returned by inspect_attachments."],
    sheet_name: Annotated[str | None, "Optional worksheet name for .xlsx files."] = None,
) -> str:
    """Profile a tabular attachment to understand schema, types, nulls, and numeric summary."""
    frame = _load_dataframe(attachment_id, sheet_name=sheet_name)
    numeric_columns = frame.select_dtypes(include=[np.number]).columns.tolist()
    missing_counts = {str(column): int(value) for column, value in frame.isna().sum().to_dict().items()}
    dtypes = {str(column): str(dtype) for column, dtype in frame.dtypes.to_dict().items()}
    distinct_counts = {
        str(column): int(frame[column].nunique(dropna=True))
        for column in frame.columns[:20]
    }
    numeric_summary = (
        frame[numeric_columns].describe().transpose().round(4).replace({np.nan: None}).to_dict(orient="index")
        if numeric_columns
        else {}
    )
    return json.dumps(
        {
            "attachment_id": attachment_id,
            "row_count": int(len(frame)),
            "column_count": int(len(frame.columns)),
            "columns": list(map(str, frame.columns.tolist())),
            "dtypes": dtypes,
            "missing_counts": missing_counts,
            "distinct_counts": distinct_counts,
            "numeric_summary": numeric_summary,
        },
        ensure_ascii=False,
        indent=2,
        default=str,
    )


def _build_repl() -> PythonREPL:
    context = get_tool_runtime_context()
    workspace_dir = context.workspace_dir.resolve()
    artifact_dir = context.artifact_dir.resolve()

    def attachment_path_fn(attachment_id: str) -> str:
        return resolve_runtime_attachment(attachment_id).storage_path

    def load_dataframe_fn(attachment_id: str, sheet_name: str | None = None) -> pd.DataFrame:
        return _load_dataframe(attachment_id, sheet_name=sheet_name)

    def read_document_fn(attachment_id: str, max_chars: int = 6000) -> str:
        return _extract_document_text_internal(attachment_id, max_chars=max_chars)["text_excerpt"]

    def artifact_path_fn(file_name: str) -> str:
        return str(artifact_path(file_name))

    def register_artifact_fn(file_name: str, title: str | None = None) -> str:
        artifact = register_runtime_artifact(file_path=file_name, title=title)
        return json.dumps(
            {
                "file_name": artifact.file_name,
                "storage_path": artifact.storage_path,
                "mime_type": artifact.mime_type,
            },
            ensure_ascii=False,
        )

    globals_dict: dict[str, Any] = {
        "__name__": "__main__",
        "pd": pd,
        "np": np,
        "plt": plt,
        "sns": sns,
        "duckdb": duckdb,
        "Path": Path,
        "json": json,
        "os": os,
        "ATTACHMENTS": attachment_manifest(),
        "WORKSPACE_DIR": str(workspace_dir),
        "ARTIFACT_DIR": str(artifact_dir),
        "attachment_path": attachment_path_fn,
        "load_dataframe": load_dataframe_fn,
        "read_document_text": read_document_fn,
        "artifact_path": artifact_path_fn,
        "register_artifact": register_artifact_fn,
    }
    return PythonREPL(_globals=globals_dict, _locals={})


def _stage_attachments_for_repl() -> None:
    context = get_tool_runtime_context()

    def _materialize(source: Path, destination: Path) -> None:
        destination.parent.mkdir(parents=True, exist_ok=True)
        if destination.exists():
            return
        try:
            destination.symlink_to(source)
        except OSError:
            shutil.copy2(source, destination)

    for attachment in attachment_manifest():
        source = Path(str(attachment["storage_path"]))
        if not source.exists() or not source.is_file():
            continue
        aliases = {source.name}
        file_name = attachment.get("file_name")
        if isinstance(file_name, str) and file_name:
            aliases.add(Path(file_name).name)

        for alias in aliases:
            _materialize(source, context.artifact_dir / alias)

        storage_path_value = str(attachment.get("storage_path") or "")
        if storage_path_value and not Path(storage_path_value).is_absolute():
            _materialize(source, context.artifact_dir / storage_path_value)


@tool
def python_repl_data_tool(
    code: Annotated[str, "Python code for data processing or visualization."],
) -> str:
    """Execute Python in a sandboxed analysis workspace for dataframe / chart tasks.

    Available libraries (pre-imported as ``_`` aliases): ``matplotlib.pyplot``
    as ``_plt``, ``matplotlib.figure.Figure`` as ``_Figure``,
    ``matplotlib.font_manager`` as ``_font_manager``, ``pandas``,
    ``numpy``, ``duckdb``. The current working directory is the per-turn
    artifact directory so saved files end up in the right place.

    File access — IMPORTANT:
    - Every attached file is **automatically symlinked into the working
      directory under its original `file_name`** (e.g. ``trend.csv``,
      ``products.json``, ``multi_sheet.xlsx``). Read attachments via the
      short ``file_name``, e.g. ``pd.read_csv("trend.csv")`` or
      ``pd.read_excel("multi_sheet.xlsx", sheet_name="sales")``.
    - **Do NOT paste absolute storage paths** like
      ``/app/apps/backend/data/uploads/csv/<uuid>.csv``. UUIDs visible in
      earlier turns may be stale; the only stable handle is the short
      ``file_name`` (or ``attachment_path("<attachment_id>")`` helper).

    Chart saving:
    - Save with a short relative name (``plt.savefig("revenue.png")``) or
      with the ``artifact_path("name.png")`` helper. The collector picks up
      every new file in the workspace and registers it automatically.

    Restrictions:
    - Network access (``socket``, ``urllib.request``, ``requests``) is
      monkey-patched to raise immediately — this REPL is for local analysis,
      not external API calls.
    - Long stdout is truncated; charts must be saved via ``_plt.savefig`` so
      the artifact-collector can pick them up. Newly created files in the
      workspace are reported back to the agent.

    Returns:
    - JSON-serialised dict with ``stdout``, ``generated_files``, and any
      ``registered_artifacts`` since the previous invocation.

    Example::

        python_repl_data_tool(code='''
            import pandas as pd
            df = pd.read_csv("trend.csv")    # short name, not absolute path
            df.groupby("month")["revenue"].sum().plot.bar()
            _plt.savefig("revenue_by_month.png")
        ''')
    """
    context = get_tool_runtime_context()
    _stage_attachments_for_repl()
    before_files = _snapshot_workspace_files(context)
    repl = _build_repl()
    prelude = f"""
import os
import socket
import urllib.request
import requests
from pathlib import Path as _Path
import matplotlib.pyplot as _plt
from matplotlib.figure import Figure as _Figure
from matplotlib import font_manager as _font_manager
os.chdir(r"{context.artifact_dir}")

def _disabled_network(*args, **kwargs):
    raise RuntimeError("Network access is disabled inside python_repl_data_tool.")

urllib.request.urlopen = _disabled_network
requests.get = _disabled_network
requests.post = _disabled_network
requests.put = _disabled_network
requests.delete = _disabled_network
requests.sessions.Session.request = _disabled_network
socket.create_connection = _disabled_network

_artifact_dir = r"{context.artifact_dir}"
_original_pyplot_savefig = _plt.savefig
_original_figure_savefig = _Figure.savefig

def _safe_pyplot_savefig(fname=None, *args, _original=_original_pyplot_savefig, _artifact_dir_value=_artifact_dir, **kwargs):
    rewritten = None if fname is None else os.path.join(_artifact_dir_value, os.path.basename(str(fname)))
    return _original(rewritten, *args, **kwargs)

def _safe_figure_savefig(self, fname=None, *args, _original=_original_figure_savefig, _artifact_dir_value=_artifact_dir, **kwargs):
    rewritten = None if fname is None else os.path.join(_artifact_dir_value, os.path.basename(str(fname)))
    return _original(self, rewritten, *args, **kwargs)

_plt.savefig = _safe_pyplot_savefig
_Figure.savefig = _safe_figure_savefig

# Reset any matplotlib figure state left over from a previous turn in the
# same backend process. Without this, accumulated figures can interact in
# subtle ways (silent savefig no-op, double-rendered axes) when the python
# REPL is invoked back-to-back across different threads.
try:
    _plt.close('all')
except Exception:
    pass

_preferred_fonts = [
    'Noto Sans CJK KR',
    'Noto Sans CJK JP',
    'Noto Sans CJK SC',
    'NanumGothic',
    'Malgun Gothic',
    'AppleGothic',
]
_available_fonts = {{font.name for font in _font_manager.fontManager.ttflist}}
for _font_name in _preferred_fonts:
    if _font_name in _available_fonts:
        _plt.rcParams['font.family'] = _font_name
        break
_plt.rcParams['axes.unicode_minus'] = False
"""
    normalized_code = code
    for legacy_prefix in (
        "/mnt/data/artifact_workspace",
        "/mnt/data",
        "/app/apps/backend/artifacts",
        "/app/artifacts",
        "/app/outputs",
    ):
        normalized_code = normalized_code.replace(legacy_prefix, str(context.artifact_dir))
    result = repl.run(f"{prelude}\n{normalized_code}")
    after_files = _snapshot_workspace_files(context)
    new_files = sorted(after_files - before_files)
    auto_registered: list[str] = []
    for file_name in new_files:
        artifact = register_runtime_artifact(file_path=file_name)
        auto_registered.append(artifact.file_name)

    return json.dumps(
        {
            "status": "success" if "Failed to execute" not in result else "error",
            "stdout": result,
            "generated_files": new_files,
            "registered_artifacts": auto_registered,
            "artifact_count": len(collect_runtime_artifacts()),
        },
        ensure_ascii=False,
        indent=2,
    )


@tool
def register_analysis_artifact(
    file_name: Annotated[str, "A file name relative to the current artifact workspace."],
    title: Annotated[str | None, "Optional title for the artifact."] = None,
) -> str:
    """Register an analysis output file so it can be attached to the assistant response."""
    try:
        artifact = register_runtime_artifact(file_path=file_name, title=title)
        status = "registered"
    except ValueError:
        existing_artifacts = collect_runtime_artifacts()
        if not existing_artifacts:
            raise
        artifact = existing_artifacts[-1]
        status = "registered_existing"
    return json.dumps(
        {
            "status": status,
            "file_name": artifact.file_name,
            "mime_type": artifact.mime_type,
            "storage_path": artifact.storage_path,
        },
        ensure_ascii=False,
    )
